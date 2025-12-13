"""
FinEdge Document Processor v2.0
Universal document loader with structure-aware chunking and table extraction.
Supports: PDF, DOCX, XLSX
"""

import os
import json
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field, asdict
from pathlib import Path
import base64


@dataclass
class TableData:
    """Represents an extracted table."""
    table_id: str
    page_number: int
    headers: List[str]
    rows: List[List[str]]
    caption: Optional[str] = None
    markdown: str = ""
    bbox: Optional[Tuple[float, float, float, float]] = None  # (x0, y0, x1, y1) for highlighting
    
    def to_dict(self) -> Dict:
        result = asdict(self)
        # Convert bbox tuple to list for JSON serialization
        if self.bbox:
            result['bbox'] = list(self.bbox)
        return result


@dataclass
class DocumentChunk:
    """A chunk of document content with rich metadata."""
    chunk_id: str
    content: str
    page_number: int
    section_title: Optional[str] = None
    chunk_type: str = "text"  # text, table, header
    tables: List[TableData] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict:
        return {
            "chunk_id": self.chunk_id,
            "content": self.content,
            "page_number": self.page_number,
            "section_title": self.section_title,
            "chunk_type": self.chunk_type,
            "tables": [t.to_dict() for t in self.tables],
            "metadata": self.metadata
        }


@dataclass
class ProcessedDocument:
    """Complete processed document with all extracted data."""
    doc_id: str
    filename: str
    file_type: str
    total_pages: int
    chunks: List[DocumentChunk]
    tables: List[TableData]
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict:
        return {
            "doc_id": self.doc_id,
            "filename": self.filename,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "chunks": [c.to_dict() for c in self.chunks],
            "tables": [t.to_dict() for t in self.tables],
            "metadata": self.metadata
        }


class DocumentProcessor:
    """
    Universal document processor for financial documents.
    Handles PDF, DOCX, and XLSX with structure-aware chunking.
    """
    
    CHUNK_SIZE = 1500
    CHUNK_OVERLAP = 200
    
    def __init__(self, storage_path: str = "data"):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)
    
    def process_file(self, file_path: str) -> ProcessedDocument:
        """Process any supported document file."""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext == ".pdf":
            return self._process_pdf(path)
        elif ext in [".docx", ".doc"]:
            return self._process_docx(path)
        elif ext in [".xlsx", ".xls"]:
            return self._process_xlsx(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID based on file content."""
        with open(file_path, "rb") as f:
            content_hash = hashlib.md5(f.read()).hexdigest()[:12]
        return f"{file_path.stem}_{content_hash}"
    
    def _process_pdf(self, file_path: Path) -> ProcessedDocument:
        """Process PDF with page tracking and enhanced table extraction using pdfplumber."""
        from langchain_community.document_loaders import PyMuPDFLoader
        
        loader = PyMuPDFLoader(str(file_path))
        pages = loader.load()
        
        doc_id = self._generate_doc_id(file_path)
        chunks = []
        all_tables = []
        
        # Extract tables with pdfplumber for accurate detection
        pdf_tables = self._extract_tables_with_pdfplumber(str(file_path), doc_id)
        all_tables.extend(pdf_tables)
        
        # Extract tables/charts from embedded images using OCR
        ocr_tables, ocr_text_chunks = self._extract_ocr_content_from_pdf(str(file_path), doc_id)
        all_tables.extend(ocr_tables)
        
        # Create a mapping of page -> tables for that page
        page_tables_map = {}
        for table in pdf_tables:
            if table.page_number not in page_tables_map:
                page_tables_map[table.page_number] = []
            page_tables_map[table.page_number].append(table)
        
        for page_idx, page in enumerate(pages):
            page_num = page_idx + 1
            page_content = page.page_content
            
            # Get tables for this page
            page_tables = page_tables_map.get(page_num, [])
            
            # Chunk the text content
            page_chunks = self._chunk_text(
                page_content, 
                page_num, 
                doc_id,
                page_tables
            )
            chunks.extend(page_chunks)
        
        # Add OCR-extracted text as additional chunks
        chunks.extend(ocr_text_chunks)
        
        return ProcessedDocument(
            doc_id=doc_id,
            filename=file_path.name,
            file_type="pdf",
            total_pages=len(pages),
            chunks=chunks,
            tables=all_tables,
            metadata={
                "source": str(file_path),
                "file_path": str(file_path),  # Store for PDF viewer
                "total_chunks": len(chunks),
                "total_tables": len(all_tables),
                "ocr_tables_count": len(ocr_tables),
                "pdfplumber_tables_count": len(pdf_tables)
            }
        )
    
    def _process_docx(self, file_path: Path) -> ProcessedDocument:
        """Process Word document."""
        from docx import Document
        
        doc = Document(str(file_path))
        doc_id = self._generate_doc_id(file_path)
        
        chunks = []
        all_tables = []
        current_section = None
        page_estimate = 1
        
        # Process paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                current_section = para.text
            full_text.append(para.text)
        
        combined_text = "\n".join(full_text)
        
        # Estimate pages (rough: 3000 chars per page)
        page_estimate = max(1, len(combined_text) // 3000)
        
        # Extract tables
        for idx, table in enumerate(doc.tables):
            table_data = self._extract_docx_table(table, idx, 1, doc_id)
            all_tables.append(table_data)
        
        # Chunk the text
        chunks = self._chunk_text(combined_text, 1, doc_id, all_tables)
        
        return ProcessedDocument(
            doc_id=doc_id,
            filename=file_path.name,
            file_type="docx",
            total_pages=page_estimate,
            chunks=chunks,
            tables=all_tables,
            metadata={
                "source": str(file_path),
                "total_chunks": len(chunks),
                "total_tables": len(all_tables)
            }
        )
    
    def _process_xlsx(self, file_path: Path) -> ProcessedDocument:
        """Process Excel spreadsheet."""
        import openpyxl
        
        wb = openpyxl.load_workbook(str(file_path), data_only=True)
        doc_id = self._generate_doc_id(file_path)
        
        chunks = []
        all_tables = []
        
        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            sheet = wb[sheet_name]
            
            # Each sheet becomes a table
            headers = []
            rows = []
            
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if row_idx == 0:
                    headers = row_data
                else:
                    if any(cell.strip() for cell in row_data):  # Skip empty rows
                        rows.append(row_data)
            
            if headers or rows:
                table_id = f"{doc_id}_sheet_{sheet_idx}"
                markdown = self._table_to_markdown(headers, rows)
                
                table_data = TableData(
                    table_id=table_id,
                    page_number=sheet_idx + 1,
                    headers=headers,
                    rows=rows,
                    caption=f"Sheet: {sheet_name}",
                    markdown=markdown
                )
                all_tables.append(table_data)
                
                # Create chunk for each sheet
                chunk = DocumentChunk(
                    chunk_id=f"{doc_id}_chunk_sheet_{sheet_idx}",
                    content=f"Sheet: {sheet_name}\n\n{markdown}",
                    page_number=sheet_idx + 1,
                    section_title=sheet_name,
                    chunk_type="table",
                    tables=[table_data],
                    metadata={"sheet_name": sheet_name}
                )
                chunks.append(chunk)
        
        return ProcessedDocument(
            doc_id=doc_id,
            filename=file_path.name,
            file_type="xlsx",
            total_pages=len(wb.sheetnames),
            chunks=chunks,
            tables=all_tables,
            metadata={
                "source": str(file_path),
                "sheet_names": wb.sheetnames,
                "total_chunks": len(chunks),
                "total_tables": len(all_tables)
            }
        )
    
    def _extract_tables_with_pdfplumber(
        self, 
        file_path: str, 
        doc_id: str
    ) -> List[TableData]:
        """
        Extract tables from PDF using pdfplumber for accurate detection.
        Returns tables with bounding box coordinates for highlighting.
        """
        tables = []
        
        try:
            import pdfplumber
        except ImportError:
            print("Warning: pdfplumber not installed. Falling back to heuristic extraction.")
            return tables
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Find all tables on this page
                    page_tables = page.find_tables()
                    
                    for table_idx, table in enumerate(page_tables):
                        # Extract table data
                        extracted = table.extract()
                        
                        if not extracted or len(extracted) < 1:
                            continue
                        
                        # First row as headers, rest as data rows
                        headers = [str(cell) if cell else "" for cell in extracted[0]]
                        rows = [
                            [str(cell) if cell else "" for cell in row]
                            for row in extracted[1:]
                        ]
                        
                        # Get bounding box for highlighting
                        bbox = table.bbox  # (x0, y0, x1, y1)
                        
                        # Create table ID
                        table_id = f"{doc_id}_table_p{page_num}_{table_idx}"
                        
                        # Convert to markdown
                        markdown = self._table_to_markdown(headers, rows)
                        
                        table_data = TableData(
                            table_id=table_id,
                            page_number=page_num,
                            headers=headers,
                            rows=rows,
                            caption=f"Table {table_idx + 1} from page {page_num}",
                            markdown=markdown,
                            bbox=bbox
                        )
                        tables.append(table_data)
                        
        except Exception as e:
            print(f"Error extracting tables with pdfplumber: {e}")
            # Return empty list, will fall back to heuristic if needed
        
        return tables
    
    def _extract_ocr_content_from_pdf(
        self,
        file_path: str,
        doc_id: str
    ) -> Tuple[List[TableData], List[DocumentChunk]]:
        """
        Extract tables and text from embedded images in PDF using OCR.
        Uses pytesseract for OCR and PyMuPDF (fitz) for image extraction.
        
        Returns:
            Tuple of (tables extracted from images, text chunks from images)
        """
        tables = []
        text_chunks = []
        
        try:
            import pytesseract
            from PIL import Image
            import fitz  # PyMuPDF
            import io
            
            # Configure Tesseract path for Windows
            import platform
            if platform.system() == 'Windows':
                import os
                # Common Tesseract installation paths on Windows
                tesseract_paths = [
                    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                    r'C:\Tesseract-OCR\tesseract.exe',
                    os.path.expanduser(r'~\AppData\Local\Tesseract-OCR\tesseract.exe'),
                ]
                for path in tesseract_paths:
                    if os.path.exists(path):
                        pytesseract.pytesseract.tesseract_cmd = path
                        break
        except ImportError as e:
            print(f"OCR dependencies not available: {e}")
            print("Install with: pip install pytesseract Pillow PyMuPDF")
            return tables, text_chunks
        
        try:
            pdf_document = fitz.open(file_path)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                image_list = page.get_images(full=True)
                
                for img_idx, img in enumerate(image_list):
                    try:
                        # Extract image
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Convert to PIL Image
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        # Skip very small images (likely icons/logos)
                        if image.width < 100 or image.height < 50:
                            continue
                        
                        # Perform OCR
                        try:
                            ocr_text = pytesseract.image_to_string(image)
                        except Exception as ocr_err:
                            print(f"OCR failed for image {img_idx} on page {page_num + 1}: {ocr_err}")
                            continue
                        
                        if not ocr_text or len(ocr_text.strip()) < 20:
                            continue  # Skip images with little text
                        
                        # Check if OCR text looks like a table
                        lines = [l for l in ocr_text.split('\n') if l.strip()]
                        is_table = self._check_if_table_text(lines)
                        
                        if is_table and len(lines) >= 2:
                            # Try to parse as table
                            table = self._parse_ocr_table(
                                lines,
                                page_num + 1,
                                len(tables),
                                doc_id,
                                "ocr_image"
                            )
                            if table:
                                tables.append(table)
                        else:
                            # Add as text chunk
                            chunk = DocumentChunk(
                                chunk_id=f"{doc_id}_ocr_p{page_num + 1}_{img_idx}",
                                content=f"[OCR from image on page {page_num + 1}]\n{ocr_text.strip()}",
                                page_number=page_num + 1,
                                section_title=f"OCR Image {img_idx + 1}",
                                chunk_type="ocr_text",
                                metadata={
                                    "source": "ocr",
                                    "image_index": img_idx,
                                    "doc_id": doc_id
                                }
                            )
                            text_chunks.append(chunk)
                            
                    except Exception as img_err:
                        print(f"Error processing image {img_idx} on page {page_num + 1}: {img_err}")
                        continue
            
            pdf_document.close()
            print(f"OCR extraction complete: {len(tables)} tables, {len(text_chunks)} text chunks from images")
            
        except Exception as e:
            print(f"Error in OCR extraction: {e}")
        
        return tables, text_chunks
    
    def _check_if_table_text(self, lines: List[str]) -> bool:
        """Check if OCR text lines look like a table structure."""
        if len(lines) < 2:
            return False
        
        # Count lines with multiple separated values (tabs, multiple spaces, pipes)
        table_like_lines = 0
        for line in lines:
            # Check for tab separators or multiple spaces or pipe characters
            if '\t' in line or '|' in line or '  ' in line:
                table_like_lines += 1
        
        # If more than 50% of lines look table-like, it's probably a table
        return table_like_lines >= len(lines) * 0.5
    
    def _parse_ocr_table(
        self,
        lines: List[str],
        page_num: int,
        table_idx: int,
        doc_id: str,
        source: str
    ) -> Optional[TableData]:
        """Parse OCR text lines into a TableData structure."""
        try:
            # Try to split lines into columns
            rows = []
            for line in lines:
                # Split by tabs, pipes, or multiple spaces
                if '\t' in line:
                    cells = [c.strip() for c in line.split('\t') if c.strip()]
                elif '|' in line:
                    cells = [c.strip() for c in line.split('|') if c.strip()]
                else:
                    # Split by 2+ spaces
                    import re
                    cells = [c.strip() for c in re.split(r'\s{2,}', line) if c.strip()]
                
                if cells:
                    rows.append(cells)
            
            if len(rows) < 2:
                return None
            
            # First row as headers
            headers = rows[0]
            data_rows = rows[1:]
            
            # Create markdown
            markdown = self._table_to_markdown(headers, data_rows)
            
            table_id = f"{doc_id}_ocr_table_p{page_num}_{table_idx}"
            
            return TableData(
                table_id=table_id,
                page_number=page_num,
                headers=headers,
                rows=data_rows,
                caption=f"OCR Table {table_idx + 1} from page {page_num}",
                markdown=markdown,
                bbox=None  # No bbox available for OCR tables
            )
            
        except Exception as e:
            print(f"Error parsing OCR table: {e}")
            return None
    
    def _extract_tables_from_text(
        self, 
        text: str, 
        page_num: int, 
        doc_id: str
    ) -> List[TableData]:
        """Extract tables from text content (heuristic-based fallback)."""
        tables = []
        lines = text.split('\n')
        
        # Simple heuristic: look for lines with consistent separators
        # This is a basic implementation; can be enhanced with ML
        current_table_lines = []
        in_table = False
        
        for line in lines:
            # Check if line looks like a table row
            delimiters = line.count('|') + line.count('\t')
            if delimiters >= 2:
                in_table = True
                current_table_lines.append(line)
            elif in_table and current_table_lines:
                # End of table
                if len(current_table_lines) >= 2:
                    table = self._parse_table_lines(
                        current_table_lines, 
                        page_num, 
                        len(tables), 
                        doc_id
                    )
                    if table:
                        tables.append(table)
                current_table_lines = []
                in_table = False
        
        return tables
    
    def _parse_table_lines(
        self, 
        lines: List[str], 
        page_num: int, 
        table_idx: int, 
        doc_id: str
    ) -> Optional[TableData]:
        """Parse table lines into structured TableData."""
        if len(lines) < 2:
            return None
        
        # Determine delimiter
        delimiter = '\t' if '\t' in lines[0] else '|'
        
        headers = [cell.strip() for cell in lines[0].split(delimiter) if cell.strip()]
        rows = []
        
        for line in lines[1:]:
            if '---' in line:  # Skip separator lines
                continue
            row = [cell.strip() for cell in line.split(delimiter) if cell.strip()]
            if row:
                rows.append(row)
        
        if not headers and not rows:
            return None
        
        table_id = f"{doc_id}_table_p{page_num}_{table_idx}"
        markdown = self._table_to_markdown(headers, rows)
        
        return TableData(
            table_id=table_id,
            page_number=page_num,
            headers=headers,
            rows=rows,
            markdown=markdown
        )
    
    def _extract_docx_table(
        self, 
        table, 
        table_idx: int, 
        page_num: int, 
        doc_id: str
    ) -> TableData:
        """Extract table from docx Table object."""
        headers = []
        rows = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if row_idx == 0:
                headers = row_data
            else:
                rows.append(row_data)
        
        table_id = f"{doc_id}_table_{table_idx}"
        markdown = self._table_to_markdown(headers, rows)
        
        return TableData(
            table_id=table_id,
            page_number=page_num,
            headers=headers,
            rows=rows,
            markdown=markdown
        )
    
    def _table_to_markdown(self, headers: List[str], rows: List[List[str]]) -> str:
        """Convert table data to markdown format."""
        if not headers and not rows:
            return ""
        
        lines = []
        
        # Headers
        if headers:
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        # Rows
        for row in rows:
            # Pad row to match header length
            padded = row + [""] * (len(headers) - len(row)) if headers else row
            lines.append("| " + " | ".join(padded[:len(headers)] if headers else padded) + " |")
        
        return "\n".join(lines)
    
    def _chunk_text(
        self, 
        text: str, 
        page_num: int, 
        doc_id: str,
        tables: List[TableData]
    ) -> List[DocumentChunk]:
        """Split text into overlapping chunks with metadata."""
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.CHUNK_SIZE,
            chunk_overlap=self.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        text_chunks = splitter.split_text(text)
        chunks = []
        
        for idx, chunk_text in enumerate(text_chunks):
            chunk = DocumentChunk(
                chunk_id=f"{doc_id}_chunk_p{page_num}_{idx}",
                content=chunk_text,
                page_number=page_num,
                chunk_type="text",
                metadata={
                    "char_count": len(chunk_text),
                    "has_tables": len(tables) > 0
                }
            )
            chunks.append(chunk)
        
        # Add tables as separate chunks for better retrieval
        for table in tables:
            table_chunk = DocumentChunk(
                chunk_id=f"{table.table_id}_chunk",
                content=f"Table:\n{table.markdown}",
                page_number=table.page_number,
                chunk_type="table",
                tables=[table],
                metadata={"is_table": True}
            )
            chunks.append(table_chunk)
        
        return chunks
    
    def get_page_content(self, file_path: str, page_number: int) -> str:
        """Get raw content for a specific page (for citation preview)."""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext == ".pdf":
            from langchain_community.document_loaders import PyMuPDFLoader
            loader = PyMuPDFLoader(str(file_path))
            pages = loader.load()
            if 0 < page_number <= len(pages):
                return pages[page_number - 1].page_content
        
        return ""


# Singleton instance
processor = DocumentProcessor()


def process_document(file_path: str) -> ProcessedDocument:
    """Convenience function to process a document."""
    return processor.process_file(file_path)
